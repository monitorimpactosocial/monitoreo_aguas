from __future__ import annotations

import json
import re
from collections import Counter
from pathlib import Path

import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"
ANEXOS_DIR = ROOT / "Anexos"
REPORT_PATH = ROOT / "Informe análisis estadístico calidad de agua PARACEL v04.docx"


KEYWORDS = [
    "paraguay",
    "rio paraguay",
    "río paraguay",
    "hierro",
    "kruskal",
    "wallis",
    "shapiro",
    "anova",
    "tukey",
    "punto",
    "entrada",
    "salida",
    "medio",
    "2021",
    "2023",
    "2025",
]


def clean_text(value: object) -> str:
    if value is None:
        return ""
    text = str(value).replace("\xa0", " ").strip()
    text = re.sub(r"\s+", " ", text)
    return text


def compact_table_sample(df: pd.DataFrame, max_rows: int = 12, max_cols: int = 10) -> list[list[str]]:
    sample = df.iloc[:max_rows, :max_cols].fillna("")
    return [[clean_text(cell) for cell in row] for row in sample.to_numpy().tolist()]


def workbook_inventory(path: Path) -> dict:
    item: dict = {
        "path": str(path.relative_to(ROOT)),
        "size_bytes": path.stat().st_size,
        "sheets": [],
        "error": None,
    }
    try:
        xls = pd.ExcelFile(path)
        for sheet_name in xls.sheet_names:
            sheet_info: dict = {"name": sheet_name, "rows": None, "cols": None, "sample": [], "keyword_hits": {}}
            try:
                df = pd.read_excel(path, sheet_name=sheet_name, header=None, dtype=object)
                df = df.dropna(how="all").dropna(axis=1, how="all")
                sheet_info["rows"] = int(df.shape[0])
                sheet_info["cols"] = int(df.shape[1])
                sheet_info["sample"] = compact_table_sample(df)
                text = " ".join(clean_text(cell).lower() for cell in df.to_numpy().flatten() if clean_text(cell))
                sheet_info["keyword_hits"] = {
                    kw: text.count(kw) for kw in KEYWORDS if text.count(kw)
                }
            except Exception as exc:  # noqa: BLE001 - inventory should continue
                sheet_info["error"] = repr(exc)
            item["sheets"].append(sheet_info)
    except Exception as exc:  # noqa: BLE001 - inventory should continue
        item["error"] = repr(exc)
    return item


def report_extract(path: Path) -> dict:
    if not path.exists():
        return {"path": str(path.relative_to(ROOT)), "error": "not found"}

    document = Document(path)
    paragraphs = [clean_text(p.text) for p in document.paragraphs if clean_text(p.text)]
    table_rows: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            line = " | ".join(cell for cell in cells if cell)
            if line:
                table_rows.append(line)

    searchable_lines = paragraphs + table_rows
    hits = []
    for line in searchable_lines:
        lower = line.lower()
        if any(kw in lower for kw in KEYWORDS):
            hits.append(line)

    return {
        "path": str(path.relative_to(ROOT)),
        "paragraph_count": len(paragraphs),
        "table_count": len(document.tables),
        "keyword_lines": hits[:240],
        "first_lines": paragraphs[:40],
    }


def summarize_inventory(inventory: list[dict], report: dict) -> dict:
    file_count = len(inventory)
    sheet_count = sum(len(item.get("sheets", [])) for item in inventory)
    keyword_counter: Counter[str] = Counter()
    candidate_sources = []

    for item in inventory:
        for sheet in item.get("sheets", []):
            for key, count in sheet.get("keyword_hits", {}).items():
                keyword_counter[key] += count
            hits = sheet.get("keyword_hits", {})
            if any(k in hits for k in ["rio paraguay", "río paraguay", "hierro", "punto", "2021", "2023"]):
                candidate_sources.append(
                    {
                        "file": item["path"],
                        "sheet": sheet["name"],
                        "rows": sheet["rows"],
                        "cols": sheet["cols"],
                        "keyword_hits": hits,
                    }
                )

    return {
        "file_count": file_count,
        "sheet_count": sheet_count,
        "keyword_totals": dict(keyword_counter.most_common()),
        "candidate_sources": candidate_sources,
        "report_keyword_line_count": len(report.get("keyword_lines", [])),
    }


def main() -> None:
    DATA_DIR.mkdir(exist_ok=True)
    excel_files = sorted(
        [p for p in ANEXOS_DIR.rglob("*") if p.suffix.lower() in {".xlsx", ".xls"}],
        key=lambda p: str(p).lower(),
    )
    inventory = [workbook_inventory(path) for path in excel_files]
    report = report_extract(REPORT_PATH)
    summary = summarize_inventory(inventory, report)

    (DATA_DIR / "source_inventory.json").write_text(
        json.dumps({"summary": summary, "workbooks": inventory, "report": report}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    (DATA_DIR / "report_extract.txt").write_text(
        "\n".join(report.get("keyword_lines", [])),
        encoding="utf-8",
    )
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
